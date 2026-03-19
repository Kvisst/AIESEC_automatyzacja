from docx import Document


def wypelnianie(sciezka_wejsciowa, sciezka_wyjsciowa, dane):
    dokument = Document(sciezka_wejsciowa)

    def zamiana_paragraf(paragraphs, dane):
        for paragraph in paragraphs:
            for key, value in dane.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

    zamiana_paragraf(dokument.paragraphs, dane)

    for table in dokument.tables:
        for row in table.rows:
            for cell in row.cells:
                zamiana_paragraf(cell.paragraphs, dane)

    dokument.save(sciezka_wyjsciowa)

if __name__ == '__main__':
    lista_osob = ["Jan", "Jakub", "Ola", "Zosia"]
    suma = str(len(lista_osob))

    agenda = ["I. Otwarcie Obrad VI Walnego Zebrania.",
    "II. Zatwierdzenie protokołu z obrad V Walnego Zebrania z dnia 01.08.2025.",
    "III. Głosowanie nad wnioskiem o nadanie członkostwa zwyczajnego kandydatom.",
    "IV. Zamknięcie obrad zebrania."]



    lista_formatowana = "\n".join(lista_osob)
    agenda_formatowana = "\n".join(agenda)

    dane = {
        '{numer}': 'V',
        '{data}': '19.03.2026',
        '{przewodniczący}': 'Zosia',
        '{lista}': lista_formatowana,
        '{protokolant}': 'Maciek',
        '{liczba}': suma,
        '{agenda}': agenda_formatowana,
    }

    wypelnianie('test.docx', 'test1.docx', dane)